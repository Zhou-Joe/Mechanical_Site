#coding:utf-8

import accdata
import matplotlib
import MplPlot
matplotlib.use('Qt5Agg')
from PyQt5 import QtCore, QtWidgets, QtGui
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg, NavigationToolbar2QT as NavigationToolbar
from matplotlib.figure import Figure
import numpy as np
import docx
from docx.shared import Inches
import sys
import os
from constants import DEFAULT_CUTOFF_FREQUENCY, MAX_CUTOFF_FREQUENCY

# Get the directory where this script is located
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

colors = ['#FF0000','#0000FF','#ADD8E6','#800080','#FFFF00', '#00FF00', '#FF00FF', '#FFC0CB', '#808080',
              '#000000', '#FFA500', '#A52A2A', '#800000', '#008000', '#808000','#00008B', '#7FFFD4', '#C0C0C0','#00FFFF']

colors_code= ['#FF0000','#0000FF','#ADD8E6','#800080','#FFFF00', '#00FF00', '#FF00FF', '#FFC0CB', '#808080',
              '#000000', '#FFA500', '#A52A2A', '#800000', '#008000', '#808000','#00008B', '#7FFFD4', '#C0C0C0','#00FFFF']


class truncateTimeWindow(QtWidgets.QDialog):
    def __init__(self, mainui, data):
        super(truncateTimeWindow, self).__init__()
        overalllayout = QtWidgets.QGridLayout()
        self.setWindowTitle('Truncate Data')
        lb_starttime = QtWidgets.QLabel('Start Time: ', self)
        lb_endtime = QtWidgets.QLabel('End Time: ', self)
        self.input_starttime = QtWidgets.QSpinBox()
        self.input_endtime = QtWidgets.QSpinBox()
        self.input_starttime.setRange(0,10000)
        self.input_endtime.setRange(0, 10000)
        ok_button = QtWidgets.QPushButton('OK', self)
        ok_button.clicked.connect(lambda: self.ok_bt_func(mainui=mainui, data=data))
        cancel_button = QtWidgets.QPushButton('Cancel', self)
        cancel_button.clicked.connect(lambda: self.cancel_bt_func())
        overalllayout.addWidget(lb_starttime, 0, 0)
        overalllayout.addWidget(lb_endtime, 1, 0)
        overalllayout.addWidget(self.input_starttime, 0, 1)
        overalllayout.addWidget(self.input_endtime, 1, 1)
        overalllayout.addWidget(ok_button, 0, 2)
        overalllayout.addWidget(cancel_button, 1, 2)
        self.setLayout(overalllayout)

    def ok_bt_func(self, mainui, data):
        try:
            data.truncate_data(starttime=self.input_starttime.value(), endtime=self.input_endtime.value(),cutoff=mainui.cutoff)
            mainui.trendplot.initplotxyz()
            mainui.clearlayout(mainui.statlayout)
            mainui.replot_all()
            self.close()

        except Exception as e:
            pass

    def cancel_bt_func(self):
        self.close()

class inputExportFilename(QtWidgets.QDialog):
    def __init__(self, mainui, data):
        super(inputExportFilename, self).__init__()
        overalllayout = QtWidgets.QGridLayout()
        self.setWindowTitle('Export')


        self.lb_inputfilename = QtWidgets.QLabel('Input Filename: ', self)
        self.input_filename = QtWidgets.QLineEdit()
        self.suffix = QtWidgets.QComboBox(self)
        self.suffix.addItems(['.sup', '.sups', '.txt'])
        self.select_export = QtWidgets.QComboBox(self)
        self.select_export.addItems(['Formated Data for Newtonviewer','Filtered Clean Data'])

        ok_button = QtWidgets.QPushButton('OK', self)
        ok_button.clicked.connect(lambda: self.ok_bt_func(mainui=mainui, data=data))
        cancel_button = QtWidgets.QPushButton('Cancel', self)
        cancel_button.clicked.connect(lambda: self.cancel_bt_func())


        overalllayout.addWidget(self.select_export, 0, 0, 1, 2)

        overalllayout.addWidget(self.lb_inputfilename, 1, 0, 1, 2)
        overalllayout.addWidget(self.input_filename, 2, 0)
        overalllayout.addWidget(self.suffix, 2,1)

        overalllayout.addWidget(ok_button, 3, 0)
        overalllayout.addWidget(cancel_button, 3, 1)
        self.setLayout(overalllayout)

    def ok_bt_func(self, mainui, data):
        try:
            if self.select_export.currentText()=='Filtered Data':
                data.export_filter(plottype=mainui.plottype, path=self.input_filename.text()+'_filtered.txt')

            else:

                data.export_to_file(path=self.input_filename.text() + self.suffix.currentText())

            self.close()

        except Exception as e:
            pass

    def cancel_bt_func(self):
        self.close()


class inputEditData(QtWidgets.QWidget):
    def __init__(self, mainui, data):
        super(inputEditData, self).__init__()
        overalllayout = QtWidgets.QGridLayout()
        self.setWindowTitle('Edit Data')
        lb_add = QtWidgets.QLabel('Add a value to: ', self)
        lb_multiply = QtWidgets.QLabel('Multiply a value to ')
        lb_addx = QtWidgets.QLabel('Add to Fore/Back (x)')
        lb_addy = QtWidgets.QLabel('Add to Left/Right (y)')
        lb_addz = QtWidgets.QLabel('Add to Up/Down (z)')

        self.input_addx=QtWidgets.QDoubleSpinBox()
        self.input_addy = QtWidgets.QDoubleSpinBox()
        self.input_addz = QtWidgets.QDoubleSpinBox()

        lb_mulx = QtWidgets.QLabel('Multiply to Fore/Back (x)')
        lb_muly = QtWidgets.QLabel('Multiply to Left/Right (y)')
        lb_mulz = QtWidgets.QLabel('Multiply to Up/Down (z)')

        self.input_mulx = QtWidgets.QDoubleSpinBox()
        self.input_muly = QtWidgets.QDoubleSpinBox()
        self.input_mulz = QtWidgets.QDoubleSpinBox()

        for spinbox in [self.input_addz, self.input_addy, self.input_addx, self.input_mulx, self.input_muly, self.input_mulz]:
            spinbox.setRange(-10, 10)
            if spinbox in [self.input_mulx, self.input_muly, self.input_mulz]:
                spinbox.setValue(1)


        ok_button1 = QtWidgets.QPushButton('Apply Add', self)
        ok_button1.clicked.connect(lambda: self.ok_bt_func1(mainui=mainui, data=data))

        ok_button2 = QtWidgets.QPushButton('Apply Multiply', self)
        ok_button2.clicked.connect(lambda: self.ok_bt_func2(mainui=mainui, data=data))

        cancel_button = QtWidgets.QPushButton('Close', self)
        cancel_button.clicked.connect(lambda: self.cancel_bt_func())


        overalllayout.addWidget(lb_add, 0, 0)
        overalllayout.addWidget(lb_addx, 1, 0)
        overalllayout.addWidget(lb_addy, 2, 0)
        overalllayout.addWidget(lb_addz, 3, 0)
        overalllayout.addWidget(self.input_addx, 1, 1)
        overalllayout.addWidget(self.input_addy, 2, 1)
        overalllayout.addWidget(self.input_addz, 3, 1)

        overalllayout.addWidget(lb_multiply, 4, 0)
        overalllayout.addWidget(lb_mulx, 5, 0)
        overalllayout.addWidget(lb_muly, 6, 0)
        overalllayout.addWidget(lb_mulz, 7, 0)

        overalllayout.addWidget(self.input_mulx, 5, 1)
        overalllayout.addWidget(self.input_muly, 6, 1)
        overalllayout.addWidget(self.input_mulz, 7, 1)



        overalllayout.addWidget(ok_button1, 0, 2)
        overalllayout.addWidget(ok_button2, 4, 2)
        overalllayout.addWidget(cancel_button, 7, 2)
        self.setLayout(overalllayout)

    def ok_bt_func1(self, mainui, data):
        try:
            data.edit_data(method='add', value_array=[self.input_addx.value(), self.input_addy.value(), self.input_addz.value()],cutoff=mainui.cutoff)
            mainui.trendplot.initplotxyz()
            mainui.clearlayout(mainui.statlayout)
            mainui.replot_all()
        except Exception as e:
            print (e)

    def ok_bt_func2(self, mainui, data):
        try:
            data.edit_data(method='multiply',
                           value_array=[self.input_mulx.value(), self.input_muly.value(), self.input_mulz.value()], cutoff=mainui.cutoff)
            mainui.trendplot.initplotxyz()
            mainui.clearlayout(mainui.statlayout)

            mainui.replot_all()
        except Exception as e:
            print (e)

    def cancel_bt_func(self):
        self.close()










class inputAngleWindow(QtWidgets.QDialog):
    def __init__(self, data, mainui):
        super(inputAngleWindow, self).__init__()
        overalllayout = QtWidgets.QGridLayout()
        self.setWindowTitle('Edit Angle')
        lb_pitch_angle = QtWidgets.QLabel('Pitch Angle (°): ', self)
        lb_seatback_angle = QtWidgets.QLabel('Seatback Angle (°): ', self)
        lb_roll_angle = QtWidgets.QLabel('Roll Angle (°): ', self)
        lb_yaw_angle = QtWidgets.QLabel('Yaw Angle (°): ', self)

        self.input_pitch = QtWidgets.QSpinBox()
        self.input_seatback = QtWidgets.QSpinBox()
        self.input_roll = QtWidgets.QSpinBox()
        self.input_yaw = QtWidgets.QSpinBox()
        for spinbox in [self.input_pitch, self.input_roll, self.input_seatback, self.input_yaw]:
            spinbox.setRange(-180, 180)
        ok_button = QtWidgets.QPushButton('OK', self)
        ok_button.clicked.connect(lambda: self.ok_bt_func(data, mainui))
        cancel_button = QtWidgets.QPushButton('Cancel', self)
        cancel_button.clicked.connect(lambda: self.cancel_bt_func())
        overalllayout.addWidget(lb_pitch_angle, 0, 0)
        overalllayout.addWidget(lb_seatback_angle, 1, 0)
        overalllayout.addWidget(lb_roll_angle, 2, 0)
        overalllayout.addWidget(lb_yaw_angle, 3, 0)
        overalllayout.addWidget(self.input_pitch, 0, 1)
        overalllayout.addWidget(self.input_seatback, 1, 1)
        overalllayout.addWidget(self.input_roll, 2, 1)
        overalllayout.addWidget(self.input_yaw, 3, 1)
        overalllayout.addWidget(ok_button, 0, 2)
        overalllayout.addWidget(cancel_button, 1,2)

        self.setLayout(overalllayout)










    def ok_bt_func(self, data, mainui):
        try:
            data.reformat(overwrite = True, setting_angle=True, pitch_angle=int(self.input_pitch.value()),
                          seatback_angle=int(self.input_seatback.value()), roll_angle=int(self.input_roll.value()),
                          yaw_angle=int(self.input_yaw.value()),cutoff = mainui.cutoff)
            mainui.trendplot.initplotxyz()
            mainui.clearlayout(mainui.statlayout)

            mainui.replot_all()



            self.close()

        except Exception as e:
            pass

    def cancel_bt_func(self):
        self.close()

class inputcutoff(QtWidgets.QDialog):
    def __init__(self, mainui):
        super(inputcutoff, self).__init__()
        overalllayout = QtWidgets.QGridLayout()
        self.setWindowTitle('Cutoff Frequency')
        lb_pitch_angle = QtWidgets.QLabel('Set New Cutoff Frequency (Hz): ', self)

        self.input_pitch = QtWidgets.QSpinBox()
        self.input_pitch.setMaximum(MAX_CUTOFF_FREQUENCY)
        self.input_pitch.setMinimum(1)
        self.input_pitch.setValue(mainui.cutoff)
        ok_button = QtWidgets.QPushButton('OK', self)
        ok_button.clicked.connect(lambda: self.ok_bt_func(mainui))
        cancel_button = QtWidgets.QPushButton('Cancel', self)
        cancel_button.clicked.connect(lambda: self.cancel_bt_func())
        overalllayout.addWidget(lb_pitch_angle, 0, 0)

        overalllayout.addWidget(self.input_pitch, 0, 1)

        overalllayout.addWidget(ok_button, 1, 0)
        overalllayout.addWidget(cancel_button, 1, 1)

        self.setLayout(overalllayout)


    def ok_bt_func(self, mainui):
        try:
            mainui.set_cutoff(int(self.input_pitch.value()))
           # print (mainui.cutoff)

            self.close()
            mainui.label_type.setText("\tCurrent Plot Data Type: {}, Current Cutoff Frequency: {}Hz.".format(mainui.plottype, mainui.cutoff))
            mainui.trendplot.initplotxyz()
            mainui.clearlayout(mainui.statlayout)
            mainui.replot_all()
        except Exception as e:
            print (e)

    def cancel_bt_func(self):
        self.close()

class popoutmssg(QtWidgets.QDialog):
    def __init__(self):
        super(popoutmssg, self).__init__()
        overalllayout = QtWidgets.QGridLayout()
        self.setWindowTitle('Message')
        lb = QtWidgets.QLabel('Export Successfully!', self)
        overalllayout.addWidget(lb, 0, 0)
        ok_button = QtWidgets.QPushButton('OK', self)
        ok_button.clicked.connect(lambda: self.cancel_bt_func())
        overalllayout.addWidget(ok_button, 1, 0)
        self.setLayout(overalllayout)
    def cancel_bt_func(self):
        self.close()

class gen_report_window(QtWidgets.QDialog):
    def __init__(self, mainui):
        super(gen_report_window, self).__init__()
        overalllayout = QtWidgets.QGridLayout()
        self.setWindowTitle('Generate Report')
        lb = QtWidgets.QLabel('Select Acceleration Record to Generate: ', self)
        lb1 = QtWidgets.QLabel('Please specify restraint type in advance for ASTM Standard Fit and make sure to use filtered and zero-adjusted data. ', self)
        #self.selector = QtWidgets.QComboBox(self)


        #self.selector.addItems(mainui.datadict.keys())
        self.checkboxes = []
        for i in mainui.datadict.keys():
            self.checkboxes.append(QtWidgets.QCheckBox(i, self))



        ok_button = QtWidgets.QPushButton('OK', self)
        ok_button.clicked.connect(lambda: self.ok_bt_func(mainui))
        cancel_button = QtWidgets.QPushButton('Cancel', self)
        cancel_button.clicked.connect(lambda: self.cancel_bt_func())
        overalllayout.addWidget(lb, 0, 0, 1, 2)

        #overalllayout.addWidget(self.selector, 1, 0)
        overalllayout.addWidget(lb1, 1, 0, 1, 2)



        for i in range(len(self.checkboxes)):
            overalllayout.addWidget(self.checkboxes[i], 2+i,0, 1,2)

        overalllayout.addWidget(ok_button, len(self.checkboxes)+2, 0)
        overalllayout.addWidget(cancel_button, len(self.checkboxes)+2, 1)

        self.setLayout(overalllayout)


    def ok_bt_func(self, mainui):
        data_list = []
        self.create_doc()
        try:
            for i in self.checkboxes:
                if i.isChecked():
                    data_list.append(i.text())

            #data = mainui.datadict[self.selector.currentText()]
            for data in data_list:
                mainui.plotASTM_checkbox(canvas=mainui.canvasASTM, name=data)
                mainui.plotGB_checkbox(canvas=mainui.canvasGB, name=data)
                mainui.plotGB_checkbox(canvas=mainui.canvasAccZone, name=data)

                mainui.replot_single(data)
                i = 0
                for canvas in [mainui.canvasASTM, mainui.canvasGB, mainui.canvasAccZone, mainui.trendplot]:
                    i+=1
                    try:

                        canvas.print_figure('img{}.jpg'.format(i), dpi=200, orientation="landscape", bbox_inches='tight')
                    except Exception as e:
                        print (e)


                mainui.clearlayout(mainui.statlayout)
                mainui.replot_all()
                self.to_word(mainui.datadict[data], data, self.doc)
            self.save_doc()
            self.close()



        except Exception as e:
            print (e)



    def cancel_bt_func(self):
        self.close()

    def create_doc(self):
        self.doc = docx.Document()

    def save_doc(self):
        self.doc.save('{}.docx'.format('acceleration_plot'))


    def to_word(self, data, name, doc):
        key_info = data.get_data_stats(data.std_data)

        p = doc.add_heading('Acceleration File Name: {}'.format(name, 2))
        h1 = doc.add_heading('Trend Plot:', 4)
        doc.add_picture('img4.jpg', width=Inches(6))

        t1 = doc.add_table(rows=1, cols=6, style="Table Grid")
        t1.autofit = True
        hdr_cells = t1.rows[0].cells
        hdr_cells[0].text = ""
        hdr_cells[1].text = key_info.columns[0]
        hdr_cells[2].text = key_info.columns[1]
        hdr_cells[3].text = key_info.columns[2]
        hdr_cells[4].text = key_info.columns[3]
        hdr_cells[5].text = key_info.columns[4]



        for i in range(key_info.shape[0]):

            row_cells = t1.add_row().cells
            row_cells[0].text = key_info.index[i]

            for j in range(key_info.shape[1]):

                row_cells[j+1].text = str(key_info.iloc[i, j])

        h2 = doc.add_heading('ASTM F2291 Fitting:', 4)
        doc.add_picture('img1.jpg', width=Inches(6))
        h3 = doc.add_heading('GB 8408 Fitting:', 4)
        doc.add_picture('img2.jpg', width=Inches(6))
        h4 = doc.add_heading('Acceleration Zone Plot:', 4)
        doc.add_picture('img3.jpg', width=Inches(6))


        for i in range(4):
            os.remove('img{}.jpg'.format(i+1))








class exportWindow(QtWidgets.QDialog):
    def __init__(self, data):
        super(exportWindow, self).__init__()
        overalllayout = QtWidgets.QGridLayout()



class infobar_content(QtWidgets.QWidget):
    def __init__(self, data, plottype, statlayout, plot_order, func_remove, mainui):
        label = QtWidgets.QLabel()
        label.setMaximumHeight(58)
        label.setStyleSheet("color: black;font-size:10.5px;background-color:white;border: 2px solid {};border-radius: 4px;padding: 2px;".format(colors_code[plot_order-1]));
        label.setText(data.get_angle_info())
        label.setAutoFillBackground(True)
        label.setMaximumWidth(350)
        label.setWordWrap(True)
        statlayout.addWidget(label)
        table = QtWidgets.QTableView()
        table.setFixedHeight(80)
        table.setFixedWidth(350)
        if plottype == "Standard":
            model = DataStatusTable(data.get_data_stats(data.std_data))
        elif plottype == "Raw":
            model = DataStatusTable(data.get_data_stats(data.data))
        elif plottype == "Filter":
            model = DataStatusTable(data.get_data_stats(data.filtered_data))
        table.setModel(model)

        table.resizeColumnsToContents()
        table.resizeRowsToContents()
        statlayout.addWidget(table)
        statlayout.setAlignment(QtCore.Qt.AlignTop)

        btlayout = QtWidgets.QHBoxLayout()
        btlayout.setSpacing(1)

        bt_export = QtWidgets.QPushButton("Export")
        bt_export.clicked.connect(lambda: self.popupInputFilename(data=data, mainui=mainui))
        bt_edit = QtWidgets.QPushButton("Edit Angle")
        bt_edit.clicked.connect(lambda: self.popupInputAngle(data, mainui))
        bt_remove = QtWidgets.QPushButton("Remove")
        bt_remove.clicked.connect(func_remove)
        bt_truncate = QtWidgets.QPushButton('Truncate')
        bt_truncate.clicked.connect(lambda: self.popupInputTime(data=data, mainui=mainui))
        bt_manipulate = QtWidgets.QPushButton('Modify')
        bt_manipulate.clicked.connect(lambda: self.popupEditData(data=data, mainui=mainui))


        for bt in [bt_truncate, bt_edit, bt_export, bt_remove, bt_manipulate]:
            bt.setStyleSheet("QPushButton{font-size:10px;}")
            bt.setFixedWidth(55)
        btlayout.addWidget(bt_truncate)
        btlayout.addWidget(bt_edit)
        btlayout.addWidget(bt_manipulate)
        btlayout.addWidget(bt_remove)
        btlayout.addWidget(bt_export)



        statlayout.addLayout(btlayout)

    def popupEditData(self, data, mainui):
        try:
            self.win = inputEditData(data=data, mainui=mainui)
            self.win.show()
        except Exception as e:
            print(e)


    def popupInputFilename(self, data, mainui):
        try:
            win = inputExportFilename(data=data, mainui=mainui)
            win.show()
        except Exception as e:
            print (e)

    def popupInputAngle(self, data, mainui):
        try:
            win = inputAngleWindow(data, mainui)
            win.show()

        except Exception as e:
            pass

    def popupInputTime(self, mainui, data):
        try:
            win = truncateTimeWindow(mainui, data)
            win.show()
        except Exception as e:
            pass




class DataStatusTable(QtCore.QAbstractTableModel):
    def __init__(self, data):
        super(DataStatusTable, self).__init__()
        self._data = data

    def data(self, index, role):
        if role == QtCore.Qt.DisplayRole:
            value = self._data.iloc[index.row(), index.column()]
            # See below for the nested-list data structure.
            # .row() indexes into the outer list,
            # .column() indexes into the sub-list
            return str(value)

    def rowCount(self, index):
        # The length of the outer list.
        return self._data.shape[0]

    def columnCount(self, index):
        # The following takes the first sub-list, and returns
        # the length (only works if all rows are an equal length)
        return self._data.shape[1]

    def headerData(self, section, orientation, role):
        # section is the index of the column/row.
        if role == QtCore.Qt.DisplayRole:
            if orientation == QtCore.Qt.Horizontal:
                return str(self._data.columns[section])

            if orientation == QtCore.Qt.Vertical:
                return str(self._data.index[section])

class Ui_MainWindow(QtWidgets.QMainWindow):
    def __init__(self, *args, **kwargs):
        super(Ui_MainWindow, self).__init__(*args, **kwargs)
        self.version = '3.0'
        self.setWindowTitle('BERT Reader V{}'.format(self.version))
        self.plottype = "Standard"
        self.annot_mode = "Single"
        self.cutoff = DEFAULT_CUTOFF_FREQUENCY
        self.select_plot_tab2 = QtWidgets.QComboBox()
        self.select_plot_tab3 = QtWidgets.QComboBox()
        self.select_plot_tab4 = QtWidgets.QComboBox()
        self.select_plot_tab5 = QtWidgets.QComboBox()
        self.select_plot_tab6 = QtWidgets.QComboBox()
        self.stat_content=[]
        self.data_namelist=[]
        self.datadict=dict()
        self.setupUi()
        self.init_plotdata()



    def set_cutoff(self, cf):
        self.cutoff = cf


    def init_plotdata(self):
        self.datalist = []

    def setupUi(self):
        self.resize(1600, 900)
        self.centralwidget = QtWidgets.QWidget(self)
        self.overallLayout = QtWidgets.QHBoxLayout()
        self.centralwidget.setLayout(self.overallLayout)
        self.setCentralWidget(self.centralwidget)
        self.scrollbar = QtWidgets.QScrollArea(self)

        self.statlayout = QtWidgets.QVBoxLayout()
        self.overallLayout.addWidget(self.scrollbar)
        self.scrollbar.setFixedWidth(400)
        self.scrollbar.setWidgetResizable(True)
        widget=QtWidgets.QWidget()
        widget.setLayout(self.statlayout)
        self.scrollbar.setWidget(widget)
        self.set_toolbar()
        self.set_tabUI()


    def set_toolbar(self):
        toolbar = QtWidgets.QToolBar('Toolbar')
        self.addToolBar(toolbar)

        button_add_rawdata = QtWidgets.QAction("Add Raw Data", self)
        button_add_rawdata.triggered.connect(lambda: self.openraw())

        #button_add_stingdata = QtWidgets.QAction("Add Sting Data |", self)
        #button_add_stingdata.triggered.connect(lambda: self.opensting())

        button_add_data = QtWidgets.QAction("Add Accel Data", self)
        button_add_data.setStatusTip("Add Standard Acceleration Data")
        button_add_data.triggered.connect(lambda: self.openandplot())


        button_clear = QtWidgets.QAction("Clear all", self)
        button_clear.setStatusTip("Clear all Datasets")
        button_clear.triggered.connect(lambda: self.resettrendplot(reset_datalist=True))


        button_set_cutoff = QtWidgets.QAction("Set Cutoff Frequency", self)
        button_set_cutoff.triggered.connect(lambda: self.setcutoff_UI())




        toolbar.addAction(button_add_rawdata)
        #toolbar.addAction(button_add_stingdata)
        toolbar.addAction(button_add_data)
        toolbar.addAction(button_clear)
        toolbar.addAction(button_set_cutoff)
        self.annot_text=QtWidgets.QLabel('(Currently: {})'.format(self.annot_mode), self)

        self.annot_text.setStyleSheet("color:#ff0000;background-color: lightblue;");

        annot_switch = QtWidgets.QAction("Switch Annotation Mode", self)
        annot_switch.triggered.connect(self.switch_annot_mode)


        btn_report=QtWidgets.QAction('Generate Report', self)

        btn_report.triggered.connect(lambda: self.select_data_for_report())

        toolbar.addAction(annot_switch)
        toolbar.addAction(btn_report)

        toolbar.addWidget(self.annot_text)






    def setcutoff_UI(self):
        try:
            win = inputcutoff(self)
            win.show()


        except Exception as e:
            pass

    def select_data_for_report(self):
        try:
            win = gen_report_window(self)
            win.show()


        except Exception as e:
            print (e)





    def switch_annot_mode(self):
        if self.annot_mode == "Single":
            self.annot_mode = "Multiple"
        else:
            self.annot_mode = 'Single'

        for canvas in [self.canvasASTM,self.canvasGB, self.canvasAccZone, self.trendplot]:
            try:
                canvas.annot_mode=self.annot_mode
            except:
                pass
        self.annot_text.setText('Currently: {}'.format(self.annot_mode))



    def clearlayout(self, layout):
        while layout.count():
            item = layout.takeAt(0)
            widget = item.widget()
            if widget != None:
                widget.deleteLater()
            else:
                self.clearlayout(item.layout())

    def set_infobar(self, data, plottype, plot_order):

        info = infobar_content(statlayout=self.statlayout,data=data, plottype=plottype,
                               plot_order=plot_order,func_remove=lambda:self.remove_button_func(plot_order),
                               mainui=self)
        self.stat_content.append(info)


    def reduce_data(self, plot_order):
        self.datalist.pop(plot_order-1)
        self.data_namelist.pop(plot_order - 1)
        self.datadict = dict(zip(self.data_namelist, self.datalist))


    def replot_all(self):
        i=1
        for d in self.datalist:
            d.filter_data(cutoff=self.cutoff)
            self.addtrendplot(data=d, plottype=self.plottype, plot_order=i)
            i+=1

    def replot_single(self, name):

        d = self.datadict[name]
        i=1
        self.trendplot.initplotxyz()


        d.filter_data(cutoff=self.cutoff)
        self.addtrendplot(data=d, plottype=self.plottype, plot_order=i)


        #self.replot_all()







    def switch_to_fitmax(self, axis):
        if len(self.datalist)>1:
            self.trendplot.initplotxyz()
            self.clearlayout(self.statlayout)
            self.replot_fitmax(axis=axis)


    def replot_fitmax(self, axis):
        offset=np.array(self.get_argmax(axis=axis))-self.get_argmax(axis=axis)[0]
        i=1
        for d in self.datalist:
            self.addtrendplot_fitmax(data=d, plottype=self.plottype, plot_order=i, offset=offset[i-1])
            i+=1

    def addtrendplot_fitmax(self, data, plottype, plot_order, offset):
        try:
            color = colors[plot_order-1]
            if plottype == "Raw":
                self.trendplot.addplotxyz_fitmax(data.data, plot_order, offset, color=color)
            elif plottype == "Standard":
                self.trendplot.addplotxyz_fitmax(data.std_data, plot_order, offset, color=color)
            elif plottype == "Filter":
                self.trendplot.addplotxyz_fitmax(data.filtered_data, plot_order, offset, color=color)
            self.trendplot.draw_idle()
            self.set_infobar(data, plottype, plot_order)
        except Exception as e:
            pass


    def reorder_colors(self, plot_order):
        temp = colors.pop(plot_order-1)
        temp_code = colors_code.pop(plot_order-1)
        colors.append(temp)
        colors_code.append(temp_code)


    def remove_button_func(self, plot_order):
        self.reduce_data(plot_order=plot_order)
        self.clearlayout(self.statlayout)
        self.reorder_colors(plot_order=plot_order)

        self.stat_content=[]

        self.trendplot.initplotxyz()
        self.clearlayout(self.statlayout)
        self.replot_all()
        for combobox in [self.select_plot_tab2,self.select_plot_tab3, self.select_plot_tab4, self.select_plot_tab5, self.select_plot_tab6]:
            combobox.clear()
            combobox.addItems(self.datadict.keys())

    def set_tabUI(self):
        self.tabWidget = QtWidgets.QTabWidget()
        self.tab1 = QtWidgets.QWidget()



        self.label_type = QtWidgets.QLabel()
        self.label_type.setText("  Current Plot Data Type: {}, Current Cutoff Frequency: {}Hz.  ".format(self.plottype, self.cutoff))

        self.label_type.setStyleSheet("color:#ff0000;background-color: lightblue;");



        ##Tab3

        l3=QtWidgets.QLabel('Select Data to Plot: ', self)
        self.tab3 = QtWidgets.QWidget()
        self.GBlayout=QtWidgets.QGridLayout(self.tab3)
        self.canvasGB = MplPlot.GBPlot()
        self.canvasGB.fig.canvas.mpl_connect("button_press_event", lambda: self.canvasGB.onclick)

        mpltoolbar3 = NavigationToolbar(self.canvasGB, self.tab3)

        toolbar_tab3=QtWidgets.QToolBar('Toolbar for Plot')
        toolbar_tab3.addWidget(l3)
        gen_plot_tab3 = QtWidgets.QAction("\t\tGenerate/Update",self)
        gen_plot_tab3.triggered.connect(lambda: self.plotGB(canvas=self.canvasGB, combobox=self.select_plot_tab3))
        toolbar_tab3.addWidget(self.select_plot_tab3)
        toolbar_tab3.addAction(gen_plot_tab3)


        self.GBlayout.addWidget(toolbar_tab3)
        self.GBlayout.addWidget(self.canvasGB)
        self.GBlayout.addWidget(mpltoolbar3)
        self.tab3.setLayout(self.GBlayout)







        ##Tab 1
        button_raw = QtWidgets.QAction("Show Raw Data ", self)
        button_raw.triggered.connect(lambda: self.switch_to_plottype(plottype="Raw"))

        button_filter = QtWidgets.QAction("Filter Only ", self)
        button_filter.triggered.connect(lambda: self.switch_to_plottype(plottype="Filter"))

        button_stdformat = QtWidgets.QAction("Filter and Calibrate", self)
        button_stdformat.triggered.connect(lambda: self.switch_to_plottype(plottype="Standard"))

        button_fitmaxX = QtWidgets.QAction("Fit Max X", self)
        button_fitmaxX.triggered.connect(lambda: self.switch_to_fitmax(axis=1))

        button_fitmaxY = QtWidgets.QAction("Fit Max Y", self)
        button_fitmaxY.triggered.connect(lambda: self.switch_to_fitmax(axis=2))

        button_fitmaxZ = QtWidgets.QAction("Fit Max Z", self)
        button_fitmaxZ.triggered.connect(lambda: self.switch_to_fitmax(axis=3))


        toolbar_tab1 = QtWidgets.QToolBar('Toolbar for Plot')
        toolbar_tab1.autoFillBackground()
        toolbar_tab1.addAction(button_raw)
        toolbar_tab1.addAction(button_filter)
        toolbar_tab1.addAction(button_stdformat)
        toolbar_tab1.addAction(button_fitmaxX)
        toolbar_tab1.addAction(button_fitmaxY)
        toolbar_tab1.addAction(button_fitmaxZ)
        toolbar_tab1.addWidget(self.label_type)

        self.trendlayout = QtWidgets.QGridLayout(self.tab1)
        self.trendplot = MplPlot.MplCanvas()
        self.trendlayout.addWidget(toolbar_tab1)
        self.trendlayout.addWidget(self.trendplot)
        toolbar2 = NavigationToolbar(self.trendplot, self.tab1)
        self.trendlayout.addWidget(toolbar2)

        # Tab 4

        self.tab4 = QtWidgets.QWidget()
        l4 = QtWidgets.QLabel('Select Data to Plot: ', self)
        self.AccZonelayout = QtWidgets.QGridLayout(self.tab4)
        self.canvasAccZone = MplPlot.AccZonePlot()
        self.canvasAccZone.fig.canvas.mpl_connect("button_press_event", self.canvasAccZone.onclick)

        mpltoolbar4 = NavigationToolbar(self.canvasAccZone, self.tab4)
        toolbar_tab4 = QtWidgets.QToolBar('Toolbar for Plot')
        toolbar_tab4.addWidget(l4)
        gen_plot_tab4 = QtWidgets.QAction("Generate/Update", self)
        gen_plot_tab4.triggered.connect(lambda: self.plotGB(canvas=self.canvasAccZone, combobox=self.select_plot_tab4))

        toolbar_tab4.addWidget(self.select_plot_tab4)
        toolbar_tab4.addAction(gen_plot_tab4)
        self.AccZonelayout.addWidget(toolbar_tab4)
        self.AccZonelayout.addWidget(self.canvasAccZone)
        self.AccZonelayout.addWidget(mpltoolbar4)

        # Tab 2 - ASTM
        self.tab2 = QtWidgets.QWidget()
        l2 = QtWidgets.QLabel('Select Data to Plot:  ', self)
        self.ASTMlayout = QtWidgets.QGridLayout(self.tab2)
        self.canvasASTM = MplPlot.ASTMPlot()
        self.canvasASTM.fig.canvas.mpl_connect("button_press_event", self.canvasASTM.onclick)

        mpltoolbar2 = NavigationToolbar(self.canvasASTM, self.tab2)
        toolbar_tab2 = QtWidgets.QToolBar('Toolbar for Plot')
        toolbar_tab2.addWidget(l2)
        gen_plot_tab2 = QtWidgets.QAction("Generate/Update", self)
        gen_plot_tab2.triggered.connect(lambda: self.plotASTM(canvas=self.canvasASTM, name=self.select_plot_tab2))

        self.select_restraint = QtWidgets.QComboBox()
        self.select_restraint.addItems(['None','Individual Lower Body','Upper Body','Group Lower Body','Convenience Restraint','No Restraint'])

        self.select_cond = QtWidgets.QComboBox()
        self.select_cond.addItems(['Normal', 'E-Stop', 'Expected/Permitted Bumping'])

        self.height_input = QtWidgets.QSpinBox()
        toolbar_tab2.addWidget(l2)
        toolbar_tab2.addWidget(self.select_plot_tab2)


        toolbar_tab2.addWidget(QtWidgets.QLabel('Select Restraint Type:  ', self))
        toolbar_tab2.addWidget(self.select_restraint)
        toolbar_tab2.addWidget(QtWidgets.QLabel('Select Run Condition:  ', self))
        toolbar_tab2.addWidget(self.select_cond)
        toolbar_tab2.addWidget(QtWidgets.QLabel('Input Patron Height: ', self))
        toolbar_tab2.addWidget(self.height_input)
        toolbar_tab2.addAction(gen_plot_tab2)
        self.ASTMlayout.addWidget(toolbar_tab2)
        self.ASTMlayout.addWidget(self.canvasASTM)
        self.ASTMlayout.addWidget(mpltoolbar2)

        # Tab 5 - Z Acceleration Reversal Analysis
        self.tab5 = QtWidgets.QWidget()
        l5 = QtWidgets.QLabel('Select Data to Plot: ', self)
        self.ZReversallayout = QtWidgets.QGridLayout(self.tab5)
        self.canvasZReversal = MplPlot.ZReversalPlot()
        self.canvasZReversal.fig.canvas.mpl_connect("button_press_event", self.canvasZReversal.onclick)

        mpltoolbar5 = NavigationToolbar(self.canvasZReversal, self.tab5)
        toolbar_tab5 = QtWidgets.QToolBar('Toolbar for Plot')
        toolbar_tab5.addWidget(l5)
        gen_plot_tab5 = QtWidgets.QAction("Generate/Update", self)
        gen_plot_tab5.triggered.connect(lambda: self.plotZReversal())

        toolbar_tab5.addWidget(self.select_plot_tab5)
        toolbar_tab5.addAction(gen_plot_tab5)
        self.ZReversallayout.addWidget(toolbar_tab5)
        self.ZReversallayout.addWidget(self.canvasZReversal)
        self.ZReversallayout.addWidget(mpltoolbar5)

        # Tab 6 - X/Y Acceleration Reversal Analysis
        self.tab6 = QtWidgets.QWidget()
        l6 = QtWidgets.QLabel('Select Data to Plot: ', self)
        self.XYReversallayout = QtWidgets.QGridLayout(self.tab6)
        self.canvasXYReversal = MplPlot.XYReversalPlot()
        self.canvasXYReversal.fig.canvas.mpl_connect("button_press_event", self.canvasXYReversal.onclick)

        mpltoolbar6 = NavigationToolbar(self.canvasXYReversal, self.tab6)
        toolbar_tab6 = QtWidgets.QToolBar('Toolbar for Plot')
        toolbar_tab6.addWidget(l6)
        gen_plot_tab6 = QtWidgets.QAction("Generate/Update", self)
        gen_plot_tab6.triggered.connect(lambda: self.plotXYReversal())

        toolbar_tab6.addWidget(self.select_plot_tab6)
        toolbar_tab6.addAction(gen_plot_tab6)
        self.XYReversallayout.addWidget(toolbar_tab6)
        self.XYReversallayout.addWidget(self.canvasXYReversal)
        self.XYReversallayout.addWidget(mpltoolbar6)

        self.overallLayout.addWidget(self.tabWidget, 1)
        self.tabWidget.addTab(self.tab1, "Trend")
        self.tabWidget.addTab(self.tab2, "Fit ASTM Contour")
        self.tabWidget.addTab(self.tab3, "Fit GB Contour")
        self.tabWidget.addTab(self.tab4, "Fit Acceleration Zone")
        self.tabWidget.addTab(self.tab5, "Z Reversal Analysis")
        self.tabWidget.addTab(self.tab6, "X/Y Reversal Analysis")




    def get_argmax(self, axis):
        arg_max=[]
        if self.plottype == "Raw":
            for data in self.datalist:
                arg_max.append(data.data.iloc[:,axis].idxmax())
        if self.plottype == "Standard":
            for data in self.datalist:
                arg_max.append(data.std_data.iloc[:,axis].idxmax())
        if self.plottype == "Filter":
            for data in self.datalist:
                arg_max.append(data.filtered_data.iloc[:,axis].idxmax())

        return arg_max


    def plotGB(self, canvas, combobox):
        try:
            data = self.datadict[combobox.currentText()]
            if self.plottype == "Standard":
                canvas.addplotxyz(data=data.std_data)
            elif self.plottype == "Filter":
                canvas.addplotxyz(data=data.filtered_data)
            elif self.plottype == "Raw":
                canvas.addplotxyz(data=data.data)
        except Exception as e:
            pass

    def plotGB_checkbox(self, canvas, name):
        try:
            data = self.datadict[name]
            if self.plottype == "Standard":
                canvas.addplotxyz(data=data.std_data)
            elif self.plottype == "Filter":
                canvas.addplotxyz(data=data.filtered_data)
            elif self.plottype == "Raw":
                canvas.addplotxyz(data=data.data)
        except Exception as e:
            pass

    def plotASTM(self, canvas, name):
        try:
            data = self.datadict[name.currentText()]
            restraint = self.select_restraint.currentText()
            cond = self.select_cond.currentText()
            input_height = self.height_input.value()
            if self.plottype == "Standard":
                data=data.std_data
            elif self.plottype == "Filter":
                data=data.filtered_data
            elif self.plottype == "Raw":
                data=data.data

            canvas.addplotxyz(data=data, restraint = restraint, cond = cond, input_height = input_height)
        except Exception as e:
            pass

    def plotASTM_checkbox(self, canvas, name): #just for checkbox
        try:
            data = self.datadict[name]
            restraint = self.select_restraint.currentText()
            cond = self.select_cond.currentText()
            input_height = self.height_input.value()
            if self.plottype == "Standard":
                data=data.std_data
            elif self.plottype == "Filter":
                data=data.filtered_data
            elif self.plottype == "Raw":
                data=data.data

            canvas.addplotxyz(data=data, restraint = restraint, cond = cond, input_height = input_height)
        except Exception as e:
            pass


    def set_datalist(self, data):
        self.datalist.append(data)
        self.data_namelist.append(data.filename)
        self.datadict=dict(zip(self.data_namelist, self.datalist))
        for combobox in [self.select_plot_tab2, self.select_plot_tab3, self.select_plot_tab4, self.select_plot_tab5, self.select_plot_tab6]:
            combobox.clear()
            combobox.addItems(self.datadict.keys())







    def openandplot(self):
        options = QtWidgets.QFileDialog.Options()
        options |= QtWidgets.QFileDialog.DontUseNativeDialog
        fileName, _ = QtWidgets.QFileDialog.getOpenFileName(self.centralwidget, "Open File", SCRIPT_DIR, "All Files (*)", options=options)
        if fileName:
            try:
                data = accdata.AccData(fileName, cutoff=self.cutoff)
                self.set_datalist(data)
                self.addtrendplot(data=self.datalist[-1], plottype=self.plottype, plot_order=len(self.datalist))
            except Exception as e:
                print (e)

    def openraw(self):
        options = QtWidgets.QFileDialog.Options()
        options |= QtWidgets.QFileDialog.DontUseNativeDialog
        fileName, _ = QtWidgets.QFileDialog.getOpenFileName(self.centralwidget, "Open File", SCRIPT_DIR, "All Files (*)", options=options)
        if fileName:
            try:
                data = accdata.RawData(fileName)

                data_GB, data_ASTM = data.export_data()
                #print(data_GB)
                data_GB.reformat(overwrite=True, setting_angle=True, pitch_angle=0, seatback_angle=0, roll_angle=0,
                                 yaw_angle=0, cutoff=self.cutoff)
                data_ASTM.reformat(overwrite=True, setting_angle=True, pitch_angle=0, seatback_angle=0, roll_angle=0,
                                   yaw_angle=0,cutoff=self.cutoff)
                self.set_datalist(data_GB)

                self.set_datalist(data_ASTM)

                for i in [-2, -1]:
                    self.addtrendplot(data=self.datalist[i], plottype=self.plottype, plot_order=len(self.datalist)+i+1)

            except Exception as e:
                pass



    def opensting(self):
        options = QtWidgets.QFileDialog.Options()
        options |= QtWidgets.QFileDialog.DontUseNativeDialog
        fileName, _ = QtWidgets.QFileDialog.getOpenFileName(self.centralwidget, "Open File", SCRIPT_DIR, "All Files (*)", options=options)
        if fileName:
            try:
                data = accdata.StingData(fileName)
                if len(data.datalist)==1:


                    data_ASTM = data.export_data()

                    data_ASTM.reformat(overwrite=True, setting_angle=True, pitch_angle=0, seatback_angle=0, roll_angle=0,
                                       yaw_angle=0,cutoff=self.cutoff)

                    self.set_datalist(data_ASTM)

                    self.addtrendplot(data=self.datalist[-1], plottype=self.plottype, plot_order=len(self.datalist))
                else:
                    data_16g, data_10hz, data_5hz = data.export_data()

                    data_16g.reformat(overwrite=True, setting_angle=True, pitch_angle=0, seatback_angle=0,
                                       roll_angle=0,
                                       yaw_angle=0,cutoff=self.cutoff)

                    data_10hz.reformat(overwrite=True, setting_angle=True, pitch_angle=0, seatback_angle=0,
                                       roll_angle=0,
                                       yaw_angle=0,cutoff=self.cutoff)
                    data_5hz.reformat(overwrite=True, setting_angle=True, pitch_angle=0, seatback_angle=0,
                                       roll_angle=0,
                                       yaw_angle=0,cutoff=self.cutoff)
                    self.set_datalist(data_16g)

                    self.set_datalist(data_10hz)

                    self.set_datalist(data_5hz)

                    for i in [-3,-2, -1]:
                        self.addtrendplot(data=self.datalist[i], plottype=self.plottype,
                                          plot_order=len(self.datalist) + i + 1)


            except Exception as e:
                pass

    def plotZReversal(self):
        """Plot Z acceleration reversal analysis"""
        try:
            from sci_calculation import detect_z_acceleration_reversal
            
            data_name = self.select_plot_tab5.currentText()
            if not data_name or data_name not in self.datadict:
                return
            
            data = self.datadict[data_name]
            
            # Get the appropriate data based on plot type
            if self.plottype == "Standard":
                plot_data = data.std_data
            elif self.plottype == "Filter":
                plot_data = data.filtered_data
            elif self.plottype == "Raw":
                plot_data = data.data
            else:
                plot_data = data.std_data
            
            # Get sampling frequency
            fs = getattr(data, 'fs', 500)
            if hasattr(data, 'calculate_sampling_frequency'):
                data.calculate_sampling_frequency()
                fs = getattr(data, 'fs', 500)
            
            # Extract time and Z acceleration data
            time_data = plot_data.iloc[:, 0].values
            z_data = plot_data.iloc[:, 3].values  # Z is column 3
            
            # Detect reversals
            reversals = detect_z_acceleration_reversal(z_data, time_data, fs)
            
            # Plot the data with reversals marked
            self.canvasZReversal.plot_z_data(time_data, z_data, reversals)
            
        except Exception as e:
            print(f"Error in plotZReversal: {e}")

    def plotXYReversal(self):
        """Plot X/Y acceleration reversal analysis"""
        try:
            from sci_calculation import detect_xy_acceleration_reversal
            
            data_name = self.select_plot_tab6.currentText()
            if not data_name or data_name not in self.datadict:
                return
            
            data = self.datadict[data_name]
            
            # Get the appropriate data based on plot type
            if self.plottype == "Standard":
                plot_data = data.std_data
            elif self.plottype == "Filter":
                plot_data = data.filtered_data
            elif self.plottype == "Raw":
                plot_data = data.data
            else:
                plot_data = data.std_data
            
            # Get sampling frequency
            fs = getattr(data, 'fs', 500)
            if hasattr(data, 'calculate_sampling_frequency'):
                data.calculate_sampling_frequency()
                fs = getattr(data, 'fs', 500)
            
            # Extract time and X/Y acceleration data
            time_data = plot_data.iloc[:, 0].values
            x_data = plot_data.iloc[:, 1].values  # X is column 1
            y_data = plot_data.iloc[:, 2].values  # Y is column 2
            
            # Detect reversals
            reversals = detect_xy_acceleration_reversal(x_data, y_data, time_data, fs)
            
            # Plot the data with reversals marked
            self.canvasXYReversal.plot_xy_data(time_data, x_data, y_data, reversals)
            
        except Exception as e:
            print(f"Error in plotXYReversal: {e}")

    def switch_to_plottype(self, plottype):
        self.plottype = plottype
        self.label_type.setText("\tCurrent Plot Data Type: {}, Current Cutoff Frequency: {}Hz.".format(self.plottype, self.cutoff))



        self.trendplot.initplotxyz()
        self.clearlayout(self.statlayout)
        self.replot_all()


    def resettrendplot(self, reset_datalist=False):
        try:
            if reset_datalist == True:
                self.datalist = []
                self.data_namelist=[]
                self.datadict=dict()
                for combobox in [self.select_plot_tab2, self.select_plot_tab3, self.select_plot_tab4, self.select_plot_tab5, self.select_plot_tab6]:
                    combobox.clear()


            self.trendplot.initplotxyz()
            self.canvasASTM.initplotxyz()
            self.canvasGB.initplotxyz()
            self.canvasAccZone.initplotxyz()
            self.canvasZReversal.initplotxyz()
            self.canvasXYReversal.initplotxyz()
            self.clearlayout(self.statlayout)
            self.trendplot.draw_idle()
            self.canvasAccZone.draw_idle()
            self.canvasGB.draw_idle()
            self.canvasASTM.draw_idle()
            self.canvasZReversal.draw_idle()
            self.canvasXYReversal.draw_idle()
            
        except Exception as e:
            pass

    def addtrendplot(self, data, plottype, plot_order):
        try:
            color = colors[plot_order-1]
            if plottype == "Raw":
                self.trendplot.addplotxyz(data.data, plot_order, color=color)
            elif plottype == "Standard":
                self.trendplot.addplotxyz(data.std_data, plot_order, color=color)
            elif plottype == "Filter":
                self.trendplot.addplotxyz(data.filtered_data, plot_order, color=color)
            self.trendplot.draw_idle()
            self.set_infobar(data, plottype, plot_order)
        except Exception as e:
            pass


if __name__ == "__main__":


    appStyle = """
    QPushButton {
        background-color: #2B5DD1;
        color: #FFFFFF;
    }
    QPushButton:hover {
        background-color: green;
    }
    
    
    QComboBox {
    font-family:  "Lucida Grande", Lucida, Verdana, sans-serif;
    border: 1px solid #D3D3D3;
    border-radius: 4px;
    background: qlineargradient(x1: 0, y1: 0, x2: 0, y2: 1,stop: 0 #EEEEEE, stop: 1 #FFFFFF);
    color: #333;
    font-size: 8pt;
    padding: 4px;
    }
     
    QComboBox:on {
    background: qlineargradient(x1: 0, y1: 0, x2: 0, y2: 1,stop: 0 #D5D5D5, stop: 1 #EEEEEE);
     }
     
    QComboBox::drop-down {
    border: 1px solid #D3D3D3;
    border-radius: 8px;
     }
     
    QComboBox::down-arrow {
    padding: 0px 5px 0px 5px;
     }
     
         QComboBox::item:selected {
    background-color: lightgreen;
     }         QComboBox::item:hover {
    background-color: lightgreen;
     }
     
     

QSpinBox {
    padding-right: 15px; /* make room for the arrows */

    border-width: 3;
}

    QSpinBox {font-family:  "Lucida Grande", Lucida, Verdana, sans-serif}
    QWidget {font-family:  "Lucida Grande", Lucida, Verdana, sans-serif;}
    
        QToolBar {background-color: white;}
        QToolButton {border: 1px solid green;border-radius: 4px;padding: 2px;}
  QToolButton:hover { background-color: lightgreen;}
    

"""




    app = QtWidgets.QApplication(sys.argv)
    app.setStyle('Fusion')
    app.setStyleSheet(appStyle)
    ui = Ui_MainWindow()

    try:
        ico_path = "icon.png"
        icon = QtGui.QIcon()
        icon.addPixmap(QtGui.QPixmap(ico_path), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        ui.setWindowIcon(icon)
    except:
        pass
    ui.show()

    sys.exit(app.exec_())

