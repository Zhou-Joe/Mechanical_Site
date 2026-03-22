from __future__ import annotations

import io
import tempfile
from pathlib import Path

try:
    import docx
    import matplotlib

    matplotlib.use("Agg")
    from docx.shared import Inches
    from matplotlib import pyplot as plt
except Exception as exc:  # pragma: no cover - exercised only when optional deps are absent
    raise RuntimeError("Report generation requires `python-docx` and `matplotlib`.") from exc

from .analysis import (
    build_astm_payload,
    build_gb_payload,
    build_zone_payload,
    get_plot_frame,
)


def build_report_document(*, datasets, plot_type: str, restraint: str, condition: str, height: float, astm_view: str, gb_view: str):
    document = docx.Document()

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_root = Path(temp_dir)

        for index, (record, acc_data) in enumerate(datasets, start=1):
            trend_path = temp_root / f"trend_{index}.png"
            astm_path = temp_root / f"astm_{index}.png"
            gb_path = temp_root / f"gb_{index}.png"
            zone_path = temp_root / f"zone_{index}.png"

            _save_trend_plot(trend_path, acc_data, plot_type, record["name"])
            _save_astm_plot(astm_path, acc_data, plot_type, restraint, condition, height, astm_view)
            _save_gb_plot(gb_path, acc_data, plot_type, gb_view)
            _save_zone_plot(zone_path, acc_data, plot_type)

            document.add_heading(f"Acceleration File Name: {record['name']}", 2)
            document.add_heading("Trend Plot", 4)
            document.add_picture(str(trend_path), width=Inches(6.4))

            stats = acc_data.get_data_stats(get_plot_frame(acc_data, plot_type))
            table = document.add_table(rows=1, cols=6, style="Table Grid")
            header = table.rows[0].cells
            header[0].text = ""
            for column_index, column_name in enumerate(stats.columns, start=1):
                header[column_index].text = str(column_name)

            for row_index in range(stats.shape[0]):
                row = table.add_row().cells
                row[0].text = str(stats.index[row_index])
                for column_index in range(stats.shape[1]):
                    row[column_index + 1].text = str(stats.iloc[row_index, column_index])

            document.add_heading(f"ASTM View: {astm_view.upper()}", 4)
            document.add_picture(str(astm_path), width=Inches(6.4))

            document.add_heading(f"GB View: {gb_view.upper()}", 4)
            document.add_picture(str(gb_path), width=Inches(6.4))

            document.add_heading("Acceleration Zone Plot", 4)
            document.add_picture(str(zone_path), width=Inches(6.4))

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output


def _save_trend_plot(path: Path, acc_data, plot_type: str, title: str) -> None:
    data = get_plot_frame(acc_data, plot_type)
    figure, axes = plt.subplots(3, 1, figsize=(10, 7), sharex=True, constrained_layout=True)
    labels = [("Fore/Aft", 1, "#e74c3c"), ("Lateral", 2, "#27ae60"), ("Vertical", 3, "#2980b9")]
    for axis, (label, index, color) in zip(axes, labels):
        axis.plot(data.iloc[:, 0], data.iloc[:, index], color=color, linewidth=0.9)
        axis.set_ylabel(label)
        axis.grid(alpha=0.25)
    axes[-1].set_xlabel("Time (s)")
    figure.suptitle(title)
    figure.savefig(path, dpi=160)
    plt.close(figure)


def _save_astm_plot(path: Path, acc_data, plot_type: str, restraint: str, condition: str, height: float, view: str) -> None:
    payload = build_astm_payload(acc_data, plot_type, restraint, condition, height)
    figure, axis = plt.subplots(figsize=(8, 5), constrained_layout=True)

    if view == "ay":
        axis.plot(payload["time_axis"], payload["ay_positive"], label="Measured ay (Right)", color="#1f77b4")
        axis.plot(payload["time_axis"], payload["ay_left"], label="Measured ay (Left)", color="#2ca02c")
        limits = payload["base_limits"]["ay_positive"]
        axis.plot(limits["x"], limits["y"], color="red", linewidth=2.4, label="ASTM")
        if payload["restraint_limits"].get("ay_positive"):
            restraint_limits = payload["restraint_limits"]["ay_positive"]
            axis.plot(restraint_limits["x"], restraint_limits["y"], color="black", linewidth=1.8, label=payload["restraint"])
        axis.set_ylabel("ay")
        axis.set_ylim(0, 3.2)
    elif view == "az":
        axis.plot(payload["time_axis"], payload["az_positive"], label="Measured az (Down)", color="#1f77b4")
        axis.plot(payload["time_axis"], payload["az_negative"], label="Measured az (Up)", color="#2ca02c")
        for key in ("az_negative", "az_positive"):
            limits = payload["base_limits"][key]
            axis.plot(limits["x"], limits["y"], color="red", linewidth=2.4, label="ASTM" if key == "az_negative" else None)
        for key in ("az_negative", "az_positive"):
            if payload["restraint_limits"].get(key):
                limits = payload["restraint_limits"][key]
                axis.plot(limits["x"], limits["y"], color="black", linewidth=1.8, label=payload["restraint"] if key == "az_negative" else None)
        axis.set_ylabel("az")
    elif view in {"xy", "xz", "yz"}:
        if view == "xy":
            axis.plot(payload["raw_x"], payload["raw_y"], color="#17becf", linewidth=0.8)
            axis.plot(payload["base_contours"]["xy"]["x"], payload["base_contours"]["xy"]["y"], "r--", linewidth=2, label="ASTM (0.2s)")
            if payload["restraint_contours"].get("xy"):
                axis.plot(payload["restraint_contours"]["xy"]["x"], payload["restraint_contours"]["xy"]["y"], color="black", linewidth=1.5, label=f"{payload['restraint']} (0s)")
            if payload["restraint_contours"].get("xy_dash"):
                axis.plot(payload["restraint_contours"]["xy_dash"]["x"], payload["restraint_contours"]["xy_dash"]["y"], "k--", linewidth=1.5, label=f"{payload['restraint']} (0.2s)")
            axis.set_xlabel("Front <=> Back")
            axis.set_ylabel("Left <=> Right")
        elif view == "xz":
            axis.plot(payload["raw_x"], payload["raw_z"], color="#17becf", linewidth=0.8)
            axis.plot(payload["base_contours"]["xz"]["x"], payload["base_contours"]["xz"]["z"], "r--", linewidth=2, label="ASTM (0.2s)")
            if payload["restraint_contours"].get("xz"):
                axis.plot(payload["restraint_contours"]["xz"]["x"], payload["restraint_contours"]["xz"]["z"], color="black", linewidth=1.5, label=f"{payload['restraint']} (0s)")
            if payload["restraint_contours"].get("xz_dash"):
                axis.plot(payload["restraint_contours"]["xz_dash"]["x"], payload["restraint_contours"]["xz_dash"]["z"], "k--", linewidth=1.5, label=f"{payload['restraint']} (0.2s)")
            axis.set_xlabel("Front <=> Back")
            axis.set_ylabel("Up <=> Down")
        else:
            axis.plot(payload["raw_y"], payload["raw_z"], color="#17becf", linewidth=0.8)
            axis.plot(payload["base_contours"]["yz"]["y"], payload["base_contours"]["yz"]["z"], "r--", linewidth=2, label="ASTM (0.2s)")
            if payload["restraint_contours"].get("yz"):
                axis.plot(payload["restraint_contours"]["yz"]["y"], payload["restraint_contours"]["yz"]["z"], color="black", linewidth=1.5, label=f"{payload['restraint']} (0s)")
            if payload["restraint_contours"].get("yz_dash"):
                axis.plot(payload["restraint_contours"]["yz_dash"]["y"], payload["restraint_contours"]["yz_dash"]["z"], "k--", linewidth=1.5, label=f"{payload['restraint']} (0.2s)")
            axis.set_xlabel("Left <=> Right")
            axis.set_ylabel("Up <=> Down")
        axis.axhline(0, color="black", linewidth=1)
        axis.axvline(0, color="black", linewidth=1)
    else:
        axis.plot(payload["time_axis"], payload["ax_positive"], label="Measured ax (+)", color="#1f77b4")
        axis.plot(payload["time_axis"], payload["ax_negative"], label="Measured ax (-)", color="#2ca02c")
        for key in ("ax_negative", "ax_positive"):
            limits = payload["base_limits"][key]
            axis.plot(limits["x"], limits["y"], color="red", linewidth=2.4, label="ASTM" if key == "ax_negative" else None)
        for key in ("ax_negative", "ax_positive"):
            if payload["restraint_limits"].get(key):
                limits = payload["restraint_limits"][key]
                axis.plot(limits["x"], limits["y"], color="black", linewidth=1.8, label=payload["restraint"] if key == "ax_negative" else None)
        axis.set_ylabel("ax")

    axis.set_xlabel("dt (s)" if view in {"ax", "ay", "az"} else "")
    axis.grid(alpha=0.25)
    axis.legend(loc="best")
    axis.set_title(f"ASTM {view.upper()}")
    figure.savefig(path, dpi=160)
    plt.close(figure)


def _save_gb_plot(path: Path, acc_data, plot_type: str, view: str) -> None:
    payload = build_gb_payload(acc_data, plot_type)
    figure, axis = plt.subplots(figsize=(8, 5), constrained_layout=True)

    if view == "az":
        axis.plot(payload["time_axis"], payload["az_positive"], label="Measured az (+)", color="#1f77b4")
        axis.plot(payload["time_axis"], payload["az_negative"], label="Measured az (-)", color="#2ca02c")
        for key in ("az_positive", "az_negative"):
            limits = payload["limits"][key]
            axis.plot(limits["x"], limits["y"], color="red", linewidth=2.4, label="Allowable az" if key == "az_positive" else None)
        axis.set_ylabel("az")
    elif view == "combined":
        axis.plot(payload["raw_z"], payload["raw_y"], color="#17becf", linewidth=0.9, label="Measured Data")
        for label, color in (("dt_005", "yellow"), ("dt_010", "orange"), ("dt_020", "red")):
            line = payload["combined_limits"][label]
            axis.plot(line["x"], line["y_pos"], color=color, linewidth=2.2, label=label.replace("_", "."))
            axis.plot(line["x"], [-value for value in line["y_pos"]], color=color, linewidth=2.2)
        axis.set_xlabel("az")
        axis.set_ylabel("ay")
    else:
        axis.plot(payload["time_axis"], payload["ay_positive"], label="Measured ay (+)", color="#1f77b4")
        axis.plot(payload["time_axis"], payload["ay_negative"], label="Measured ay (-)", color="#2ca02c")
        for key in ("ay_positive", "ay_negative"):
            limits = payload["limits"][key]
            axis.plot(limits["x"], limits["y"], color="red", linewidth=2.4, label="Allowable ay" if key == "ay_positive" else None)
        axis.set_ylabel("ay")

    if view != "combined":
        axis.set_xlabel("dt (s)")
    axis.grid(alpha=0.25)
    axis.legend(loc="best")
    axis.set_title(f"GB {view.upper()}")
    figure.savefig(path, dpi=160)
    plt.close(figure)


def _save_zone_plot(path: Path, acc_data, plot_type: str) -> None:
    payload = build_zone_payload(acc_data, plot_type)
    figure, axis = plt.subplots(figsize=(8, 6), constrained_layout=True)

    for zones in payload["zone_boundaries"].values():
        for zone in zones:
            axis.fill(zone["x"], zone["y"], color=zone["color"])

    axis.scatter(payload["zone_data"]["x"], payload["zone_data"]["z"], c=payload["zone_data"]["colors"], s=8)
    axis.set_xlim(-6, 6)
    axis.set_ylim(-4, 4)
    axis.set_xlabel("Front <=> Back (X-acceleration)")
    axis.set_ylabel("Up <=> Down (Z-acceleration)")
    axis.set_title("Acceleration Zones")
    axis.grid(alpha=0.2)

    most_severe = payload["zone_analysis"]["most_severe"]
    if most_severe:
        axis.text(
            0.02,
            0.98,
            f"Classified Zone: {most_severe['zone']}",
            transform=axis.transAxes,
            va="top",
            bbox={"boxstyle": "round", "facecolor": "wheat", "alpha": 0.65},
        )

    figure.savefig(path, dpi=160)
    plt.close(figure)
